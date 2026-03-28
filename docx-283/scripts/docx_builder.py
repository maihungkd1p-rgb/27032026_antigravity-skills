import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class DocxBuilder:
    """
    Wrapper cho python-docx giúp tạo file Word chuẩn Corporate (Antigravity V4).
    Áp dụng triết lý 4-Layers: Cấu trúc (TOC/Headings) -> Thông tin (Bảng/Lists) -> Trực quan (Styles).
    """
    
    def __init__(self, title: str):
        self.doc = Document()
        self.title = title
        self._setup_styles()
        self.add_title(title)

    def _setup_styles(self):
        """Khởi tạo các Style chuẩn Corporate (màu sắc, font Arial)"""
        styles = self.doc.styles
        
        # Sửa Base Normal Font
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        
        # Sửa Title Font
        if 'Title' in styles:
            title_style = styles['Title']
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
            
        # Sửa Heading 1 Font
        if 'Heading 1' in styles:
            h1_style = styles['Heading 1']
            h1_font = h1_style.font
            h1_font.name = 'Arial'
            h1_font.size = Pt(16)
            h1_font.bold = True
            h1_font.color.rgb = RGBColor(0, 51, 102)

        # Sửa Heading 2 Font
        if 'Heading 2' in styles:
            h2_style = styles['Heading 2']
            h2_font = h2_style.font
            h2_font.name = 'Arial'
            h2_font.size = Pt(14)
            h2_font.bold = True
            h2_font.color.rgb = RGBColor(0, 81, 152)

    def add_title(self, text: str):
        """Thêm Tiêu đề chính căn giữa"""
        p = self.doc.add_paragraph(text, style='Title')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def add_toc_placeholder(self):
        """Thêm không gian giả định cho Mục lục (TOC)"""
        p = self.doc.add_paragraph()
        run = p.add_run("[MỤC LỤC - CHÈN TỰ ĐỘNG TẠI ĐÂY SAU KHI MỞ WORD]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_page_break()

    def add_heading(self, text: str, level: int = 1):
        """Thêm Heading (Chương mục) có đánh số hoặc phân cấp rõ ràng"""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold=False, italic=False):
        """Thêm đoạn văn body text chuẩn"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        if bold: run.bold = True
        if italic: run.italic = True
        return p

    def add_bullet_list(self, items: list):
        """Thêm danh sách dạng Bullet chia nhỏ text block dài"""
        for item in items:
            self.doc.add_paragraph(str(item), style='List Bullet')

    def add_numbered_list(self, items: list):
        """Thêm danh sách dạng Numbering quy trình"""
        for item in items:
            self.doc.add_paragraph(str(item), style='List Number')

    def add_table(self, headers: list, data: list):
        """Thêm bảng dữ liệu định lượng (Trực quan hóa Information)"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Medium Shading 1 Accent 1' # Chuẩn Corporate table có header xanh
        
        # Fill Headers
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        # Fill Data
        for row_data in data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                
    def add_page_break(self):
        """Thêm ngắt trang bắt buộc giữa các Phần Lớn (Structure Layer)"""
        self.doc.add_page_break()

    def save(self, filepath: str):
        """Lưu file (Đảm bảo lưu vào 03_Outputs)"""
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        self.doc.save(filepath)
        print(f"✅ Đã lưu file Word thành công tại: {filepath}")

# ----- Usage Example -----
if __name__ == "__main__":
    doc = DocxBuilder("Báo Cáo Tình Hình Kinh Doanh Q1/2026")
    doc.add_toc_placeholder()
    
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph("Quý 1 chứng kiến sự tăng trưởng vượt bậc 15% so với cùng kỳ năm trước. Dưới đây là các ý chính:", bold=True)
    doc.add_bullet_list([
        "Doanh thu đạt 120 tỷ VNĐ.",
        "Chi phí marketing tối ưu giảm 5%.",
        "Mở mới 3 cửa hàng tại khu vực miền Nam."
    ])
    
    doc.add_page_break()
    doc.add_heading("2. Chỉ Số Tài Chính", level=1)
    doc.add_paragraph("Chi tiết các chỉ số phân bổ theo khu vực:")
    
    headers = ["Khu vực", "Doanh thu (Tỷ)", "Tăng trưởng", "Đánh giá"]
    data = [
        ["Miền Bắc", "65", "+10%", "Hoàn thành"],
        ["Miền Trung", "20", "+5%", "Chậm tiến độ"],
        ["Miền Nam", "35", "+25%", "Vượt chỉ tiêu"]
    ]
    doc.add_table(headers, data)
    
    doc.save("03_Outputs/Bao_Cao_Q1_2026_Sample.docx")
